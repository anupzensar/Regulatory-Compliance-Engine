import cv2
import playwright
import requests
from requests.exceptions import Timeout, ConnectionError
import os
import json
from dotenv import load_dotenv, get_key
import shutil
import re
from datetime import datetime
import streamlit as st
from ultralytics import YOLO
import numpy as np
from PIL import Image
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches

load_dotenv()

env_path = ".env"
installed_games = None  # Global variable to store the list of installed games
gtp_base_url = get_key(env_path, "GTP_BASE_URL")


def gtp_login():
    """Login to GTP using credentials from environment variables. Returns the access token"""
    print("Logging in to GTP...")
    st.toast("Logging in to GTP...")
    # Retrieve credentials from environment variables
    username = get_key(env_path, "GTP_USERNAME")
    password = get_key(env_path, "GTP_PASSWORD")

    if not username or not password:
        raise ValueError("Username or password is missing from environment variables.")

    # Set the login URL
    login_url = f"{gtp_base_url}games/api/v1/auth"

    # Prepare the payload (the data to be sent in the POST request)
    payload = {"username": username, "password": password}

    try:
        # Send the POST request to login
        response = requests.post(login_url, json=payload)

        # Check if the response was successful (status code 200)
        if response.status_code == 200:
            # Assuming the response contains a JSON with an access token
            response_data = response.json()
            access_token = response_data.get("token")

            if access_token:
                print("Logged in to GTP!")
                st.toast("Logged in to GTP!")
                return access_token
            else:
                raise ValueError("Access token not found in the response.")
        else:
            response.raise_for_status()

    except requests.exceptions.RequestException as e:
        raise SystemExit(f"Error during login request: {e}")


def api_request(method, url, platform="axiom", **kwargs):
    """Make an API request with authorization headers."""
    if "axiom_api_key" in st.session_state:
        axiom_api_key = st.session_state.axiom_api_key
    else:
        axiom_api_key = get_key(env_path, "AXIOM_API_KEY")
    
    if "gtp_bearer_token" not in st.session_state:
        st.session_state.gtp_bearer_token = gtp_login()

    headers = {}
    if platform == "axiom":
        headers = {"x-api-key": axiom_api_key}
    elif platform == "gtp":
        headers = {
            "Authorization": f"Bearer {st.session_state.gtp_bearer_token}",
        }

    kwargs["headers"] = headers

    retries = 3
    for attempt in range(retries):
        try:
            if method.lower() == "post":
                response = requests.post(url, **kwargs)
            elif method.lower() == "get":
                response = requests.get(url, **kwargs)
            elif method.lower() == "put":
                response = requests.put(url, **kwargs)
            elif method.lower() == "delete":
                response = requests.delete(url, **kwargs)
            elif method.lower() == "patch":
                response = requests.patch(url, **kwargs)
            else:
                raise ValueError(f"Unsupported method: {method}")

            # Check for 401 status code (Unauthorized)
            if response.status_code == 401:
                if platform == "gtp":
                    token = gtp_login()
                    set_state("gtp_bearer_token", token)

                    continue  # Retry the request
                elif platform == "axiom":
                    st.error("Request failed, please check the Axiom API key.")
                    st.stop()

            return response

        except (Timeout, ConnectionError) as e:
            print(f"Attempt {attempt + 1} failed: {e}")
            if attempt < retries - 1:
                continue
            else:
                raise


def load_json(filename):
    """Load the games configuration from a JSON file."""
    try:
        with open(filename, "r") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return None


def get_project_id(game_id: int):
    """Fetch the project ID associated with a given game ID."""
    try:
        response = api_request(
            "get", f"{gtp_base_url}api/v1/projects/games/{game_id}", "gtp"
        )

        if response.status_code == 200:
            projects = json.loads(response.content)
            # Sort the original list in place by 'createdOn' field in descending order
            projects.sort(
                key=lambda x: datetime.fromisoformat(x["createdOn"].replace("Z", "+00:00")), 
                reverse=True
            )
            for project in projects:
                name = project.get("projectName", "").lower()
                status = project.get("projectStatus", "").lower()
                if "sweep" in name and "Archived" not in status:
                    return project.get("id")
            print(f"Sweepstakes project not found for game ID: {game_id}")
            return None
        else:
            print(f"Error fetching project ID: {response.status_code}")
            return None
    except Exception as e:
        print(f"Error occurred while fetching project ID: {e}")
        return None


def get_markets(project_id, variants):
  try:
    markets = []
    payout_variants = "&payoutVariants=" + "&payoutVariants=".join([str(variant['id']) for variant in variants])
    url = f"{gtp_base_url}api/v2/attributes/projects/{project_id}?key=markets{payout_variants}"
    response = api_request("get", url, "gtp")
    
    if response.status_code == 200:
        data = response.json()
        for variant in data:
            attributes = variant.get("attributes", [])
            for attribute in attributes:
                if attribute.get("name") == "markets":
                    values = attribute.get("value", "").split(",")
                    for value in values:
                        markets.append(value.strip())
                        
        return set(markets)
    else:
        print(f"Could not retrieve markets: {response.status_code}")
        return []
  except Exception as e:
    print("Error fetching markets: ", e)
    return []


def get_state(name):
    if name in st.session_state:
        return st.session_state[name]
    else:
        print(f"Key {name} not found in session state")
        return None


def set_state(name, value):
    try:
        st.session_state[name] = value
        return True
    except Exception as e:
        print(f"Error occurred while setting session state: {e}")
        return False


def checkUser(lobby,userName):
    print("Checking userName")

    get_url = f"https://axiomcore-app1-{lobby}.installprogram.eu/UserAccounts"

    try:
        # Send the POST request with headers and JSON payload
        response = api_request("get", get_url, "axiom")

        # Check if the response was successful
        if response.status_code == 200:
            data = response.json()

            # Check if the username exists in the dataobj
            user_exists = any(user["username"] == userName for user in data.get("dataObject", []))
            if user_exists:
                print("User already Exists ",user_exists)
                return True
            else:
                return False
        else:
            response.raise_for_status()
        return False
    except requests.exceptions.RequestException as e:
        print(f"Error during request: {e}")
        return False


def createUser(user_data, userName, lobby):
    
    userExists = checkUser(lobby,userName)
   
    if ( userExists == False):

        post_url = f"https://axiomcore-app1-{lobby}.installprogram.eu/UserAccounts"
        try:
            # Send the POST request with headers and JSON payload
            response = api_request("post", post_url, "axiom", json=user_data)

            # Check if the response was successful
            if response.status_code == 200:
                print("User creation was successful!")
                st.write("User creation was successful!")
                return True
            else:
                response.raise_for_status()

        except requests.exceptions.RequestException as e:
            print(f"Error during POST request: {e}")
            st.error(f"Error during POST request: {e}")
            return False
    else:
        print(f"Error user exists {userName} ")
        st.write(f"User exists {userName}")
        return True
    return 


def get_titan_versions():
    env = get_state("env")
    if not env:
        st.error("Please enter lobby ID to fetch titan versions.")
        return []
    url = f"https://axiomcore-app1-{env}.installprogram.eu/MobileSettings/TitanVersions"
    response = api_request("get", url)
    if response.status_code == 200:
        games = response.json().get("dataObject", [])
        return (item['appVersion'] for item in games)
    return []


def get_build_versions(variant):
    env = get_state("env")
    if not env:
        st.error("Please enter lobby ID to fetch build versions.")
        return []
    url = f"https://axiomcore-app1-{env}.installprogram.eu/Games/InstalledGameRecords"
    
    response = api_request("get", url)
    if response.status_code == 200:
        data = response.json()
        games = data.get("dataObject")
        module_id = variant.get("moduleId")
        client_id = variant.get("clientId")
        for game in games:
            if int(game.get("moduleId", 0)) == int(module_id) and int(game.get("clientId", 0)) == int(client_id):
                versions = game.get("versions", [])
                return [version.get("version") for version in versions]
    else:
        print("Error fetching build versions: ", response.status_code)
        return []

# async def detect_and_click_btn(page, indices):
#     if not indices:
#         print("Indices not provided to detect button.")
#         return False
    
#     attempts = 0
#     doc = Document()
#     doc.add_heading('Button Detection Report', 0)
    
#     while attempts < 3:
#         try:
#             screenshot = await page.screenshot(full_page=True)
            
#             # Convert the screenshot (in bytes) to an OpenCV image (numpy array)
#             img = Image.open(io.BytesIO(screenshot))  # Open the image from bytes
#             img = np.array(img)  # Convert it into a numpy array (HWC format)
            
#             model = YOLO(f"./models/best copy.pt")
#             results = model(img, imgsz=640, conf=0.1, classes=indices)
            
#             for result in results:
#                 for box in result.boxes:
#                     x1, y1, x2, y2 = map(int, box.xyxy[0])  # Bounding box coordinates
#                     conf = float(box.conf[0])  # Confidence score
#                     label = result.names[int(box.cls[0])]  # Object label

#                     if conf > 0.1:  # Confidence threshold
#                         # Calculate center of the detected object
#                         click_x, click_y = (x1 + x2) // 2, (y1 + y2) // 2

#                         print(f"Detected {label} at ({click_x}, {click_y}) with confidence {conf}")
                        
#                         # Take a screenshot and save it
#                         timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
#                         screenshot_path = f"screenshot_{timestamp}.png"
#                         screenshot.save(screenshot_path)
                        
#                         # Add to the report
#                         doc.add_paragraph(f"Button Detected: {label}")
#                         doc.add_paragraph(f"Coordinates: ({click_x}, {click_y})")
#                         doc.add_paragraph(f"Confidence: {conf:.2f}")
#                         doc.add_paragraph(f"Screenshot saved as: {screenshot_path}")
                        
#                         # Add screenshot to the document
#                         doc.add_paragraph(f'Screenshot taken:')
#                         doc.add_picture(screenshot_path)
                        
#                         # Click on detected object
#                         await page.mouse.click(click_x, click_y)
                        
#                         # Save the document after a successful click
#                         report_path = f"button_detection_report_{timestamp}.docx"
#                         doc.save(report_path)
                        
#                         return True, click_x, click_y, report_path
#         except Exception as e:
#             print("Error while detecting button: ", e)
#         finally:
#             attempts += 1
    
#     print(f"Button {indices} not found.")
#     return False, None, None, None

# async def detect_and_click_btn(page, indices): 
#     if not indices:
#         print("Indices not provided to detect button.")
#         return False, None, None

#     attempts = 0
#     doc = Document()
#     doc.add_heading('Button Detection Report', 0)

#     while attempts < 3:
#         try:
#             screenshot = await page.screenshot(full_page=True)
#             img = Image.open(io.BytesIO(screenshot))
#             img_np = np.array(img)

#             model = YOLO("./models/best copy.pt")
#             results = model(img_np, imgsz=640, conf=0.1, classes=indices)

#             for result in results:
#                 for box in result.boxes:
#                     x1, y1, x2, y2 = map(int, box.xyxy[0])
#                     conf = float(box.conf[0])
#                     label = result.names[int(box.cls[0])]

#                     if conf > 0.1:
#                         click_x, click_y = (x1 + x2) // 2, (y1 + y2) // 2
#                         print(f"Detected {label} at ({click_x}, {click_y}) with confidence {conf}")

#                         await page.mouse.click(click_x, click_y)

#                         timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
#                         screenshot_path = f"screenshot_{timestamp}.png"
#                         img.save(screenshot_path)

#                         doc.add_paragraph(f"Button Detected: {label}")
#                         doc.add_paragraph(f"Coordinates: ({click_x}, {click_y})")
#                         doc.add_paragraph(f"Confidence: {conf:.2f}")
#                         doc.add_paragraph(f"Screenshot saved as: {screenshot_path}")

#                         report_path = f"button_detection_report_{timestamp}.docx"
#                         doc.save(report_path)

#                         return True, click_x, click_y
#         except Exception as e:
#             print("Error while detecting button:", e)
#         finally:
#             attempts += 1

#     print(f"Button {indices} not found.")
#     return False, None, None

# async def detect_and_click_btn(page, indices):
#     if not indices:
#         print("Indices not provided to detect button.")
#         return False
    
#     attempts = 0
#     while attempts < 3:
#         try:
#             screenshot = await page.screenshot(full_page=True)
            
#             # Convert the screenshot (in bytes) to an OpenCV image (numpy array)
#             img = Image.open(io.BytesIO(screenshot))  # Open the image from bytes
#             img = np.array(img)  # Convert it into a numpy array (HWC format)
            
            
#             model = YOLO(f"./models/best copy.pt")
#             results = model(img, imgsz=640, conf=0.1, classes=indices)
            
#             for result in results:
#                 for box in result.boxes:
#                     x1, y1, x2, y2 = map(int, box.xyxy[0])  # Bounding box coordinates
#                     conf = float(box.conf[0])  # Confidence score
#                     label = result.names[int(box.cls[0])]  # Object label

#                     if conf > 0.1:  # Confidence threshold
#                         # Calculate center of the detected object
#                         click_x, click_y = (x1 + x2) // 2, (y1 + y2) // 2

#                         print(
#                             f"Detected {label} at ({click_x}, {click_y}) with confidence {conf}"
#                         )

#                         # Click on detected object
#                         await page.mouse.click(click_x, click_y)
#                         return True, click_x, click_y
#         except Exception as e:
#             print("Error while detecting button: ", e)
#         finally:
#             attempts+=1
#     print(f"Button {indices} not found.")
#     return False, None, None

async def detect_and_click_btn(page, indices):
    if not indices:
        print("Indices not provided to detect button.")
        return False, None, None

    attempts = 0
    detected = False
    doc = Document()
    doc.add_heading('Button Detection Report', 0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = f"button_detection_report_{timestamp}.docx"

    while attempts < 3:
        try:
            screenshot_bytes = await page.screenshot(full_page=True)
            img_pil = Image.open(io.BytesIO(screenshot_bytes))
            img_np = np.array(img_pil)

            model = YOLO("./models/best copy.pt")
            results = model(img_np, imgsz=640, conf=0.1, classes=indices)

            for result in results:
                for box in result.boxes:
                    x1, y1, x2, y2 = map(int, box.xyxy[0])
                    conf = float(box.conf[0])
                    label = result.names[int(box.cls[0])]

                    if conf > 0.1:
                        click_x, click_y = (x1 + x2) // 2, (y1 + y2) // 2
                        print(f"Detected {label} at ({click_x}, {click_y}) with confidence {conf}")

                        await page.mouse.click(click_x, click_y)

                        # Add details to the doc
                        doc.add_paragraph(f"Button Detected: {label}")
                        doc.add_paragraph(f"Coordinates: ({click_x}, {click_y})")
                        doc.add_paragraph(f"Confidence: {conf:.2f}")

                        # Save and embed screenshot
                        img_crop = img_pil.crop((x1, y1, x2, y2))  # Optional: just the button
                        screenshot_buffer = io.BytesIO()
                        img_crop.save(screenshot_buffer, format='PNG')
                        screenshot_buffer.seek(0)
                        doc.add_picture(screenshot_buffer, width=Inches(2))  # Resize as needed

                        # Save and embed screenshot
                        img_crop = img_pil # Optional: just the button
                        screenshot_buffer = io.BytesIO()
                        img_crop.save(screenshot_buffer, format='PNG')
                        screenshot_buffer.seek(0)
                        doc.add_picture(screenshot_buffer, width=Inches(2))  # Resize as needed

                        doc.add_paragraph("")  # spacing
                        detected = True
        except Exception as e:
            print("Error while detecting button:", e)
        finally:
            attempts += 1

        if detected:
            break  # if something is detected, stop retrying

    if detected:
        doc.save(report_path)
        print(f"Report saved: {report_path}")
        return True, click_x, click_y

    print(f"Buttons {indices} not found.")
    return False, None, None


def sanitize_game_name(game_name):
    sanitized_name = game_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    return sanitized_name


async def take_screenshot(page, folder_name):
    screenshot_folder = f"Screenshots/{folder_name}"
    # Create a folder for saving screenshots
    if not os.path.exists(screenshot_folder):
        os.makedirs(screenshot_folder)

    # Timestamp in the screenshot filename
    formatted_timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    game_name = st.session_state.selected_game.get("displayTitle", "")
    sanitized_game_name = sanitize_game_name(game_name)
    screenshot_path = os.path.join(screenshot_folder, f"{sanitized_game_name}_{folder_name}_{formatted_timestamp}.png")

    # Capture a screenshot of the entire page
    await page.screenshot(path=screenshot_path, full_page=True)
    print(f"Screenshot captured and saved as {screenshot_path}")
    
    return screenshot_path


async def is_game_loaded(page):
    pass


def save_response_to_file(file_name, response_data, folder_name):
    """Save the raw JSON response as a text file with unique filenames."""
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    file_path = os.path.join(folder_name, file_name)


    with open(file_path, "w", encoding="utf-8") as f:
        f.write(response_data)
    print(f"Response saved to {file_path}")


def extract_file_name(url):
    """Extract the file name from the URL."""
    return url.replace("\\", "/").split("/")[-1]


def search_play_files(file_name):
    """Check if the filename starts with 'play', 'spins', or 'refresh'."""
    return file_name.lower().startswith(('play', 'spins', 'refresh'))


async def launch_site_with_playwright(p : playwright, url, status, is_headless):
    """Launch the site and manipulate network requests with Playwright."""
    status.update(label="Launching game to record transactions...", state="running")
    browser = await p.chromium.launch(headless=is_headless)
    page = await browser.new_page()

    async def handle_request(route, request):
        print(f"Intercepted request: {request.url}")
        route.continue_()

    async def handle_response(response): #data prep to store in folder
        try:
            url = response.url
            file_name = extract_file_name(url)
            if search_play_files(file_name):
                print("Extracting values from packet.")
                response_body = await response.body()  
                response_str = response_body.decode('utf-8')
                save_response_to_file(file_name, response_str, "./playFiles")
            
        except Exception as e:
            print(f"Error processing the response: {e}")
            

    page.on('route', handle_request)
    page.on('response', handle_response)

    await page.goto(url)
    page.set_default_timeout(0)

    await page.wait_for_load_state("load")
    status.update(label="Game launched!", state="running")
    # await page.wait_for_load_state("networkidle", timeout=0)
    
    return page, browser